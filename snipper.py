import os
import sys
import time
import json
import ctypes
from threading import Thread

import tkinter as tk
from tkinter import W, N, S, E, NW, LEFT
from tkinter import BooleanVar, messagebox, filedialog, ttk

import docx
from docx.shared import Inches

from PIL import Image, ImageGrab, ImageTk
from pynput import keyboard
from plyer import notification


class SnipperTool():
    def __init__(self,
                 master,
                 title='Snipper Tool',
                 width=500,
                 height=350,
                 dict_image_format={'PNG':'.png', 'BMP':'.bmp', 'JPEG':'.jpeg'},
                 dict_image_editor_application={'MS Paint': 'mspaint'}):

        self.master = master
        self.title = title
        self.width = width
        self.height = height

        self.dict_image_format = dict_image_format
        self.dict_image_editor_application = dict_image_editor_application

        self.control_padd = 10 # padding between widgets
        self.blank_col_padd = 50 # padding blank col

        self.text_box_dir_path = None
        self.btn_browse = None
        self.btn_start = None
        self.btn_create_docx = None
        self.btn_exit = None
        self.chk_box_save_setings = BooleanVar()

        self.logo = None
        self.bg_image = None
        self.path_icon = None
        self.path_settings_file = None

        self.cmb_box_image_ext = None
        self.cmb_box_image_editor = None

        self.thread_start_process = None
        self.show_image_captured_notification = True # No usage at the moment

        self._draw_window()


    # region Drawing Window

    def _draw_window(self):

        self.master.title(self.title)
        self._set_file_path()

        self._set_window_size_and_position()
        self._draw_icon_and_logo()
        self._write_header_and_instructions()
        self._draw_text_box()
        self._draw_buttons()
        self._draw_settings()

    def _set_window_size_and_position(self):
        # get screen width and height
        width_screen = self.master.winfo_screenwidth()
        height_screen = self.master.winfo_screenheight()

        # calculate position (x and y coordinates) for the Tk root window
        x_window_pos = (width_screen/2) - (self.width/2)
        y_window_pos = (height_screen/2) - (self.height/2)

        # set the dimensions & position of window
        self.master.geometry('%dx%d+%d+%d' % (self.width, self.height, x_window_pos, y_window_pos))
        # self.master.resizable(width=False, height=False)

    def _draw_icon_and_logo(self):
        # 6 is for blank label
        padding_col = self.blank_col_padd - (self.control_padd/2) - 6
        # Left blank column
        lbl_pad = tk.Label(self.master, text="", )
        lbl_pad.grid(row=0, column=0, padx=(padding_col, 0), rowspan=11, sticky=W)
        # Right blank column
        lbl_pad = tk.Label(self.master, text="", )
        lbl_pad.grid(row=0, column=7, padx=(0, padding_col), rowspan=11, sticky=W)

        if self.path_icon is not None and os.path.exists(self.path_icon):
            self.master.iconbitmap(self.path_icon)

            image_logo = (Image.open(self.path_icon)).resize((50, 50))
            self.logo = ImageTk.PhotoImage(image_logo)

            tk_label_logo = tk.Label(self.master)
            tk_label_logo.image = self.logo
            tk_label_logo['image'] = self.logo
            tk_label_logo.grid(row=0, column=1, columnspan=2,
                               padx=(self.control_padd/2, 0), pady=15, sticky=W+E)

    def _create_label(self, label_text, x_text_pos, y_text_pos, font=("Helvetica", 8)):
        '''This will draw label on window'''
        label = tk.Label(self.master, text=label_text, font=font)
        label.place(x=x_text_pos, y=y_text_pos)

    def _write_header_and_instructions(self):

        # Header
        label_header = "Snipper Tool"
        header = tk.Label(self.master, text=label_header, font=("Arial", 25))
        header.grid(row=0, column=3, columnspan=5, sticky=W)

        # Instructions
        font_inst = ("Helvetica", 8)

        label_inst_header = "Instructions :"
        label_inst_header = tk.Label(self.master, text=label_inst_header, font=font_inst,)
        label_inst_header.grid(row=7, column=1, columnspan=6, sticky=W,
                               padx=self.control_padd/2, pady=(self.control_padd/2, 0))

        label_inst = "1: Click on the ' Browse ' button to select the folder.\n" + \
                     "2: Click on the ' Start ' button to start the process.\n" + \
                     "3: Press ' PrintScreen ' key to capture the screenshot.\n" + \
                     "4: Press ' Insert ' key to capture a screenshot & edit.\n" + \
                     "5: Click on ' Build Docx ' button to create a document.\n" + \
                     "    With images available in selected folder."

        label_inst = tk.Label(self.master, text=label_inst, font=font_inst,
                              justify=LEFT, wraplength=400, anchor=NW)
        label_inst.grid(row=8, column=1, columnspan=6, padx=self.control_padd/2, sticky=W)

    def _draw_text_box(self):
        self.text_box_dir_path = tk.Entry(self.master, bd=2, width=35)
        self.text_box_dir_path.grid(row=1, column=1, columnspan=4, sticky=W+E,
                                    padx=self.control_padd/2, pady=(self.control_padd/2, 0))
        self.text_box_dir_path.delete(0, tk.END)

    def _draw_buttons(self):

        self.btn_browse = tk.Button(self.master, text="Browse", width=10, command=self._browse_dir)
        self.btn_browse.grid(row=1, column=5, sticky=W+N+E+S,
                             padx=self.control_padd/2, pady=self.control_padd/2)

        self.btn_start = tk.Button(self.master, text="Start", command=self._start_capture_process)
        self.btn_start.grid(row=3, column=1, columnspan=2, sticky=W+N+E+S,
                            padx=self.control_padd/2, pady=self.control_padd/2)

        self.btn_create_docx = tk.Button(self.master, text="Create Docx",
                                         command=self._create_document)
        self.btn_create_docx.grid(row=3, column=3, columnspan=2, sticky=W+N+E+S,
                                  padx=self.control_padd/2, pady=self.control_padd/2)

        self.btn_exit = tk.Button(self.master, text="Exit", width=10, command=self.exit)
        self.btn_exit.grid(row=3, column=5, sticky=W+N+E+S,
                           padx=self.control_padd/2, pady=self.control_padd/2)

    def _draw_settings(self):
        font_text = ('Callibri', '8')
        # width_drop_down = 
        # Image extension type drop-down
        label_type = tk.Label(self.master, text="Type :")
                                                # "Editor:"
        label_type.grid(row=2, column=1, sticky=W,
                        padx=(self.control_padd/2, 0), pady=self.control_padd/2)

        self.cmb_box_image_ext = ttk.Combobox(self.master, width=8,
                                              state='readonly', font=font_text,
                                              values=list(self.dict_image_format.keys()))
        self.cmb_box_image_ext.grid(row=2, column=2, sticky=W,
                                    # padx=(0, 0), 
                                    pady=self.control_padd/2)

        # Image editor type drop-down
        label_editor = tk.Label(self.master, text="Editor:", )
        label_editor.grid(row=2, column=3, sticky=W,
                          padx=(self.control_padd/2, 0), pady=self.control_padd/2)

        image_editors = list(self.dict_image_editor_application.keys())
        self.cmb_box_image_editor = ttk.Combobox(self.master, width=8, state='readonly',
                                                 font=font_text, values=image_editors)
        self.cmb_box_image_editor.grid(row=2, column=4, sticky=W,
                                    #    padx=(0, 0), 
                                       pady=self.control_padd/2)

        # Save config check box
        chk_box_save = tk.Checkbutton(self.master, text='Save config',
                                      variable=self.chk_box_save_setings,
                                      onvalue=True, offvalue=False, command=self._save_setting)
        chk_box_save.grid(row=2, column=5, sticky=W,
                          padx=self.control_padd/2, pady=self.control_padd/2)

        self._set_default_settings_value()

    # endregion Drawing Window


    # region Settings

    def _set_file_path(self):
        assets_path = os.path.join(os.getcwd(), 'assets')

        if not os.path.exists(assets_path):
            os.mkdir(assets_path)
        self.path_icon = os.path.join(assets_path, 'SnipperIcon.ico')

        # For executable or installer to have write access
        # Creating settings in user APPDATA folder
        path_settings = os.path.join(os.getenv('APPDATA'), self.title)
        if not os.path.exists(path_settings):
            os.mkdir(path_settings)
        self.path_settings_file = os.path.join(path_settings, 'settings.json')

    def _set_default_settings_value(self):
        self.cmb_box_image_ext.current(0)
        self.cmb_box_image_editor.current(0)
        self.chk_box_save_setings.set(False)

        self._load_settings()

        self.cmb_box_image_ext.bind("<<ComboboxSelected>>", self._save_setting)
        self.cmb_box_image_editor.bind("<<ComboboxSelected>>", self._save_setting)

    def _load_settings(self):
        if os.path.exists(self.path_settings_file):
            settings = None
            with open(self.path_settings_file, "r") as read_file:
                settings = json.load(read_file)

            list_keys = list(self.dict_image_format.keys())
            if 'image_type' in settings and settings['image_type'] in list_keys:
                self.cmb_box_image_ext.current(list_keys.index(settings['image_type']))

            list_keys = list(self.dict_image_editor_application.keys())
            if 'image_editor' in settings and settings['image_editor'] in list_keys:
                self.cmb_box_image_editor.current(list_keys.index(settings['image_editor']))

            if 'save_config' in settings and isinstance(settings['save_config'], bool):
                self.chk_box_save_setings.set(settings['save_config'])

    def _save_setting(self, event=None):
        try:
            if self.chk_box_save_setings.get():
                settings = {'image_type': self.cmb_box_image_ext.get(),
                            'image_editor': self.cmb_box_image_editor.get(),
                            'save_config': self.chk_box_save_setings.get()}
                with open(self.path_settings_file, "w") as write_file:
                    json.dump(settings, write_file)
            else:
                if os.path.exists(self.path_settings_file):
                    os.remove(self.path_settings_file)
        except Exception as error:
            messagebox.showinfo(self.title, str(error))

    # endregion Settings


    def _browse_dir(self):
        dir_path = filedialog.askdirectory()
        self.text_box_dir_path.delete(0, tk.END)
        self.text_box_dir_path.insert(0, dir_path)


    # region Capture Process

    def _start_capture_process(self):
        if os.path.exists(self.text_box_dir_path.get()) and self.thread_start_process is None:
            self.thread_start_process = Thread(target=self._listen_key_events, daemon=True)
            self.thread_start_process.start()

            messagebox.showinfo(self.title, 'Process started.')
            self.master.iconify() # Minimises window tool
        elif not os.path.exists(self.text_box_dir_path.get()):
            messagebox.showinfo(self.title, "Please select correct path.")
        elif self.thread_start_process is not None:
            self._notify('Process already running.', 3, self.title, self.path_icon)

    def _listen_key_events(self):
        with keyboard.Listener(on_release=self._on_key_relase) as listener:
            listener.join()

    def _on_key_relase(self, key):
        if key == keyboard.Key.print_screen:
            # self._capture_screen(open_image=False)
            thread = Thread(target=self._capture_screen, args=(False, ), daemon=True)
            thread.start()
        elif key == keyboard.Key.insert:
            # self._capture_screen(open_image=True)
            thread = Thread(target=self._capture_screen, args=(True, ), daemon=True)
            thread.start()

    def _capture_screen(self, open_image=False, prefix_name='Screen_Shot_'):
        image_file_extn = self.dict_image_format[self.cmb_box_image_ext.get()]
        image_file_name = prefix_name + self.time_stamp() + image_file_extn
        image_file_path = os.path.join(self.text_box_dir_path.get(), image_file_name)

        ImageGrab.grab().save(image_file_path)

        if self.show_image_captured_notification:
            message = 'Image captured ' + image_file_name
            self._notify(message, 2, self.title, self.path_icon)

        if open_image:
            app_name = self.dict_image_editor_application[self.cmb_box_image_editor.get()]
            open_image_in_editor_cmd = app_name + ' ' + image_file_path
            os.system(open_image_in_editor_cmd)

    # endregion Capture Process


    # region Create Document

    def _create_document(self):
        """Creates word document wiht images present in selected directory"""

        if not os.path.exists(self.text_box_dir_path.get()):
            messagebox.showinfo(self.title, 'Please select correct path.')
            return

        list_path_images_in_dir = self._get_list_path_images()

        # If any images in selected directory
        if list_path_images_in_dir:
            doc = docx.Document()

            # Setting page layout of 1 inch margin from all sides
            for section in doc.sections:
                section.top_margin, section.bottom_margin = Inches(1), Inches(1)
                section.left_margin, section.right_margin = Inches(1), Inches(1)

            # Adding images to word document
            for index, file_path in enumerate(list_path_images_in_dir):
                doc.add_heading(str(index + 1), 4).style = 'List'
                doc.add_picture(file_path, width=Inches(6.5), height=Inches(3.65))
                doc.add_paragraph('') # For new line

            # Saving word document
            doc_name = 'Document_' + self.time_stamp() + '.docx'
            path_doc = os.path.abspath(os.path.join(self.text_box_dir_path.get(), doc_name))
            doc.save(path_doc)

            # Show messsage
            count_images = str(len(list_path_images_in_dir))
            msg = "Word document created with " + count_images + " images at " + path_doc + '.'
            messagebox.showinfo(self.title, msg)
        else:
            messagebox.showinfo(self.title, "No images present in selected directory.")

    def _get_list_path_images(self):
        list_path_images_in_dir = []
        image_extensions = tuple(list(self.dict_image_format.values()) + ['.jpg'])

        # Finding all images present in selected directory
        for file_name in os.listdir(self.text_box_dir_path.get()):
            if file_name.lower().endswith(image_extensions):
                path_image = os.path.abspath(os.path.join(self.text_box_dir_path.get(), file_name))
                list_path_images_in_dir.append(path_image)

        return list_path_images_in_dir

    # endregion Create Document


    def exit(self):
        """To close the application"""
        self.master.destroy()


    def time_stamp(self):
        """Returns time stamp with custom format as string"""
        return str(time.strftime("%Y-%m-%d-%H-%M-%S"))


    def _notify(self, message: str, timeout: int, title: str, app_icon: str) -> None:
        """
        Shows notification message

        Args:
            message (str):  the message to display
            timeout (int):  the time for which notification should be visible
            title (str):    title of notification
            app_icon (str): accessible path of icon file
        """
        notification.notify(title=title, message=message,
                            app_name=title, app_icon=app_icon,
                            timeout=timeout)


# DPI aware
if 'win' in sys.platform:
    ctypes.windll.shcore.SetProcessDpiAwareness(1)


if __name__ == "__main__":
    MASTER = tk.Tk()
    SnipperTool(MASTER)
    MASTER.mainloop()
